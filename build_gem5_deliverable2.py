from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

OUT = str(Path(__file__).resolve().parent / "Gem5_Deliverable_2_Implementation_Guide.docx")

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.5)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
# A compact APA-style student-paper layout is used to hold the required
# implementation narrative to exactly four body pages.
normal.paragraph_format.line_spacing = 1.45
normal.paragraph_format.space_after = Pt(0)

for name, size in [("Title", 16), ("Heading 1", 11), ("Heading 2", 11)]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.paragraph_format.line_spacing = 1.45
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(0)

if "Code Block" not in styles:
    code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
else:
    code_style = styles["Code Block"]
code_style.font.name = "Courier New"
code_style._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
code_style.font.size = Pt(8.5)
code_style.paragraph_format.line_spacing = 1
code_style.paragraph_format.space_before = Pt(3)
code_style.paragraph_format.space_after = Pt(3)
code_style.paragraph_format.left_indent = Inches(.25)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer._p.append(field)

def para(text="", bold_start=None, center=False, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    if indent and not center:
        p.paragraph_format.first_line_indent = Inches(.5)
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p

def heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p

def code(lines):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    p.add_run(lines)
    return p

def page():
    doc.add_page_break()

# Title page
for _ in range(4):
    para("", center=True, indent=False)
p = para("Phase 2: Implementation Using gem5", center=True, indent=False)
p.runs[0].bold = True
p.runs[0].font.size = Pt(16)
para("Deliverable 2: Low-Power Microprocessor Implementation Guide", center=True, indent=False)
for _ in range(4):
    para("", center=True, indent=False)
for line in ["Student Name", "Institution", "Course Number and Name", "Instructor Name", "July 26, 2026"]:
    para(line, center=True, indent=False)

# Content page 1
page()
heading("Phase 2: Implementation Using gem5")
para("This guide implements a reproducible x86-64 microprocessor study in gem5 v25.1.0.1. The baseline is a detailed out-of-order core in syscall-emulation (SE) mode with a 2 GHz, 1.0 V clock domain; 32 KiB, two-way L1 instruction and data caches; a 256 KiB, eight-way shared L2 cache; 512 MiB DDR3-1600 memory; and a tournament branch predictor. Low-power operation is represented by three controlled design points that combine static voltage/frequency scaling with pipeline-width gating. The goal is architectural comparison, not transistor-level signoff. gem5 provides timing CPU models and clock/voltage domains suitable for this type of energy-efficiency exploration (gem5, 2026a).")
heading("Installation and Configuration", 2)
para("Docker Desktop provides the most reproducible installation on macOS, Windows, or Linux. Allocate at least 8 GB of memory and 25 GB of free storage; 16 GB is preferred. Apple Silicon hosts can run the supplied linux/amd64 image through emulation, although compilation and simulation are slower. From a terminal, unpack the project, enter its directory, mark the scripts executable, and build the pinned image:")
code("unzip ilp-gem5-part2.zip\ncd ilp-gem5-part2\nchmod +x scripts/*.sh\n./scripts/docker_build.sh")
para("The Dockerfile installs Ubuntu 24.04, GCC, SCons, Python development packages, protobuf, Boost, HDF5, and related libraries. It then checks out the signed gem5 v25.1.0.1 tag and builds build/X86/gem5.opt. Pinning the release prevents later changes in defaults or statistic names from silently altering the experiment (gem5 Project, 2026). Confirm the simulator and compile the statically linked workloads with:")
code("docker compose run --rm ilp gem5.opt --version\n./scripts/docker_smoke_test.sh")
para("A successful smoke test must print the benchmark checksum, a normal gem5 exit cause, and SMOKE TEST PASSED. The output directory must contain config.ini, config.json, stats.txt, and terminal.txt. The generated configuration files are the authoritative record of every instantiated SimObject and parameter (Lowe-Power, 2026).")

# Content page 2
page()
heading("Step-by-Step Implementation Process")
heading("System and Memory Hierarchy", 2)
para("First, configs/common.py constructs a System, assigns timing memory mode, and creates the root SrcClockDomain and VoltageDomain. The code attaches separate 32 KiB L1 caches to the CPU, connects both to an L2 crossbar and 256 KiB L2 cache, then connects the L2 and DDR3 controller to the system crossbar. Cache latency, associativity, MSHRs, and targets per MSHR are explicit, so the intended hierarchy can be verified in config.ini. SEWorkload.init_compatible loads a statically linked x86-64 Linux binary without booting an operating system. This choice shortens preliminary tests but excludes OS scheduling and device power.")
heading("Processor Model and Low-Power Features", 2)
para("Second, configs/run_experiment.py selects MinorCPU for pipeline tracing and DerivO3CPU for branch, issue-width, SMT, and low-power experiments. The low_power mode adds three operating profiles. Performance uses 2.0 GHz, 1.0 V, and four-wide fetch through commit. Balanced uses 1.6 GHz, 0.9 V, and two-wide operation. Eco uses 1.2 GHz, 0.8 V, and one-wide operation. The same tournament predictor, cache hierarchy, benchmark, and instruction count are retained to isolate the operating-point change.")
code('profile = {\n "performance": {"clock":"2GHz","voltage":"1.0V","width":4},\n "balanced": {"clock":"1.6GHz","voltage":"0.9V","width":2},\n "eco": {"clock":"1.2GHz","voltage":"0.8V","width":1},\n}[args.power_profile]')
para("Static DVFS is modeled by passing the selected frequency and voltage into SrcClockDomain and VoltageDomain. Width gating is modeled by setting fetchWidth, decodeWidth, renameWidth, dispatchWidth, issueWidth, wbWidth, and commitWidth to the selected width. This approximates disabling unused pipeline capacity. It does not, by itself, calculate watts or model physical clock-gating cells. A later extension can attach MathExprPowerModel expressions or McPAT; detailed timing CPUs are required for gem5’s power-model path (gem5, 2026b).")
heading("Running the Design", 2)
code("./scripts/docker_run_all.sh\n# or one profile:\ngem5.opt -d results/low_power_eco configs/run_experiment.py \\\n --experiment low_power --power-profile eco \\\n --predictor tournament --binary build/ilp_bench --arg 750000")

# Content page 3
page()
heading("Preliminary Simulation Results")
para("Preliminary validation was performed in layers. Static Python compilation succeeded for all configuration and analysis scripts; shell syntax validation succeeded for every runner; and the ILP benchmark compiled and executed natively, producing checksum 13363162888169815236 for 10,000 iterations. These checks validate syntax, benchmark determinism, and command plumbing. The first Docker build also reached the X86 gem5 C++ compilation stage, confirming dependency resolution and source checkout. Because the first build was still compiling at the reporting checkpoint, no architectural performance values are claimed here.")
heading("Required Result Collection", 2)
para("After the image finishes, run ./scripts/docker_smoke_test.sh followed by ./scripts/docker_run_all.sh. Each profile writes to a separate results/low_power_<profile> directory. Record simInsts, system.cpu.numCycles, IPC (or committed instructions divided by cycles), simSeconds, branch mispredictions, and L1/L2 demand misses from stats.txt. Verify clock and voltage in config.ini and stats.txt. gem5 statistics are hierarchical, and stats.txt may contain multiple dump blocks; use the final complete block unless the script deliberately resets statistics (Lowe-Power, 2026).")
table = doc.add_table(rows=1, cols=5)
table.autofit = False
widths = [1.35, .85, .75, 1.25, 2.05]
headers = ["Profile", "GHz", "V", "Width", "Expected validation trend"]
for i, (cell, w, text) in enumerate(zip(table.rows[0].cells, widths, headers)):
    cell.width = Inches(w)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = text
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(9)
for row in [
    ("Performance", "2.0", "1.0", "4", "Highest IPC/throughput; highest energy proxy"),
    ("Balanced", "1.6", "0.9", "2", "Moderate slowdown with lower V²f proxy"),
    ("Eco", "1.2", "0.8", "1", "Lowest throughput and lowest V²f proxy"),
]:
    cells = table.add_row().cells
    for cell, w, text in zip(cells, widths, row):
        cell.width = Inches(w)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = text
        for p in cell.paragraphs:
            p.paragraph_format.line_spacing = 1
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9)
para("Table 1. Low-power design points and hypotheses.", indent=False)
para("For a first-order comparison, compute the normalized dynamic-power proxy P* = (V/1.0 V)² × (f/2.0 GHz) × (width/4). This produces 1.000, 0.324, and 0.096 for performance, balanced, and eco, respectively. This proxy is a design-screening metric—not measured power—because activity factor, capacitance, leakage, memory energy, and execution-unit utilization are omitted. Report energy proxy as E* = P* × normalized execution time and compare it with IPC and elapsed simulated time.")

# Content page 4
page()
heading("Validation, Debugging, and Refinement")
heading("Correctness Procedure", 2)
para("Validation begins with the smoke workload and proceeds to controlled comparisons. First, confirm the benchmark output checksum is identical across all profiles; a mismatch indicates functional or argument-passing error. Second, inspect each config.ini for DerivO3CPU, timing mode, tournament predictor, cache sizes, clock, voltage, and all seven pipeline widths. Third, require simInsts > 0, numCycles > 0, a normal exit cause, and no fatal or panic text in terminal.txt. Fourth, compare committed instruction counts. Minor differences can occur around process startup, but the same workload and stopping rule should produce comparable counts. Finally, rerun one profile and confirm identical architectural statistics.")
heading("Interpretation and Acceptance Criteria", 2)
para("The implementation is accepted when all three profiles complete, configuration evidence matches the intended design, checksums agree, and statistics are internally consistent. IPC should generally fall as width is reduced, while simulated seconds should rise as frequency is reduced. However, a memory-bound workload may show little IPC benefit from the performance profile. That outcome is valid if cache-miss and memory-stall evidence explains it. The low-power claim should be framed as a performance–energy trade-off and supported by either the normalized proxy or a calibrated power model, never by frequency alone.")
heading("Debugging and Refinement", 2)
para("If gem5 reports an executable-format error, rebuild the benchmark as a static x86-64 Linux binary inside the container. If stats.txt is missing, inspect terminal.txt before rerunning. If a chosen width does not appear in config.ini, check that --experiment low_power was supplied and that the selected profile overrides the generic --width value. If voltage and frequency do not match, verify that make_system receives clock and voltage arguments before m5.instantiate(). If results are noisy or dominated by startup, increase the iteration count, fast-forward to a region of interest, reset statistics, and simulate a fixed instruction window.")
para("The next refinement is calibrated power estimation. Attach a MathExprPowerModel to the timing CPU, define dynamic and static expressions from clock period, voltage, IPC, and activity statistics, and validate coefficients against a documented source. Periodic statistic dumps can then show phase behavior, but dump frequency must be chosen carefully because it controls temporal resolution (gem5, 2026b). Full-system simulation is recommended if OS-driven DVFS, idle states, interrupts, or scheduler behavior are part of the architectural specification.")

# References page
page()
heading("References")
refs = [
    "gem5. (2026a). About gem5. https://www.gem5.org/about/",
    "gem5. (2026b). ARM power modelling. https://www.gem5.org/documentation/learning_gem5/part2/arm_power_modelling/",
    "gem5 Project. (2026). Releases: Version 25.1.0.1. GitHub. https://github.com/gem5/gem5/releases/tag/v25.1.0.1",
    "Lowe-Power, J. (2026). Understanding gem5 statistics and output. gem5. https://www.gem5.org/documentation/learning_gem5/part1/gem5_stats/",
]
for text in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(.5)
    p.paragraph_format.first_line_indent = Inches(-.5)
    p.add_run(text)

doc.core_properties.title = "Phase 2: Implementation Using gem5"
doc.core_properties.subject = "Deliverable 2 low-power microprocessor implementation guide"
doc.core_properties.author = "Student Name"
doc.save(OUT)
print(OUT)
